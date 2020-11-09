from cpa_std.vehicle_insurance.common.tools import Tool
from openpyxl import load_workbook
from docx import Document

class Record:
    def __init__(self):
        self.t=Tool()
        self.excel = f"{self.t.data}/asset/template.xlsx"
        self.res_word=f"{self.t.res}/qr.docx"
        self.res_excel=f"{self.t.res}/data.xlsx"
        self.wb=load_workbook(self.excel)
        self.ws=self.wb['测试']
        self.start=2
        self.doc=Document()
    def write_row(self,pdf,cepai,fadongji,cejia,baodanhao,begin,end,jine,toubaoren,cezu,beibaoxianren,tebie):
        self.ws.cell(row=self.start,column=2).value=cepai
        self.ws.cell(row=self.start,column=3).value=fadongji
        self.ws.cell(row=self.start,column=4).value=cejia
        self.ws.cell(row=self.start,column=5).value=baodanhao
        self.ws.cell(row=self.start,column=6).value=begin
        self.ws.cell(row=self.start,column=7).value=end
        self.ws.cell(row=self.start,column=8).value=jine
        self.ws.cell(row=self.start,column=9).value=toubaoren
        self.ws.cell(row=self.start,column=10).value=cezu
        self.ws.cell(row=self.start,column=11).value=beibaoxianren
        self.ws.cell(row=self.start,column=12).value=tebie
        self.ws.cell(row=self.start,column=13).value=pdf
        self.start=self.start+1
    def write_qr(self,path,cepai,begin,end,which,jine):
        self.doc.add_heading("二维码", 3)
    
        p =self.doc.add_paragraph(f'{cepai}--{begin}--{end}-{which}--{jine}')
        r = p.add_run()
        r.add_picture(path)
